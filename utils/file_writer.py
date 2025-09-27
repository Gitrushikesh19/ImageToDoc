from docx import Document
import os

file_name = 'output.docx'

def append_to_docx(text):
    if not os.path.exists(file_name):
        doc = Document()
        doc.add_heading('Captured Text', 0)

    else:
        doc = Document(file_name)

    doc.add_paragraph(text)
    doc.save(file_name)
